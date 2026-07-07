"""
Retrieval-Augmented Generation (RAG) pipeline.

Handles: extracting text from PDF/DOCX/TXT uploads, splitting into chunks,
generating embeddings, persisting them in a local ChromaDB vector store, and
retrieving the most relevant chunks for a given user query.
"""
import os
from typing import List

from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings

from config import config
from utils.logger import get_logger

logger = get_logger(__name__)

_embeddings = None
_vectorstore = None


def get_embeddings():
    """Lazily load the sentence-transformer embedding model (free, local, no API key)."""
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(model_name=config.EMBEDDING_MODEL)
    return _embeddings


def get_vectorstore() -> Chroma:
    """Lazily initialize the persistent Chroma vector store."""
    global _vectorstore
    if _vectorstore is None:
        _vectorstore = Chroma(
            collection_name="college_knowledge_base",
            embedding_function=get_embeddings(),
            persist_directory=config.CHROMA_PERSIST_DIR,
        )
    return _vectorstore


def extract_text(filepath: str, file_type: str) -> str:
    """Extract raw text from a PDF, DOCX, or TXT file."""
    file_type = file_type.lower()

    if file_type == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if file_type == "docx":
        import docx
        doc = docx.Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)

    if file_type == "txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    raise ValueError(f"Unsupported file type: {file_type}")


def chunk_text(text: str, source_name: str, category: str = "general") -> List[LCDocument]:
    """Split extracted text into overlapping chunks with source metadata attached."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(text)
    return [
        LCDocument(page_content=chunk, metadata={"source": source_name, "category": category})
        for chunk in chunks if chunk.strip()
    ]


def ingest_document(filepath: str, original_filename: str, category: str = "general") -> int:
    """
    Full ingestion pipeline for one uploaded document:
    extract -> chunk -> embed -> store in ChromaDB.
    Returns the number of chunks indexed.
    """
    ext = original_filename.rsplit(".", 1)[1].lower()
    text = extract_text(filepath, ext)

    if not text or not text.strip():
        raise ValueError("No extractable text found in document.")

    docs = chunk_text(text, source_name=original_filename, category=category)
    if not docs:
        raise ValueError("Document produced no valid chunks after splitting.")

    store = get_vectorstore()
    store.add_documents(docs)
    logger.info("Indexed %d chunks from '%s' into ChromaDB.", len(docs), original_filename)
    return len(docs)


def retrieve_context(query: str, top_k: int | None = None) -> tuple[str, List[dict]]:
    """
    Retrieve the most relevant document chunks for a query.
    Returns a tuple of (formatted_context_string, list_of_source_metadata_dicts).
    """
    store = get_vectorstore()
    k = top_k or config.RETRIEVAL_TOP_K

    try:
        results = store.similarity_search_with_relevance_scores(query, k=k)
    except Exception as exc:
        logger.error("Retrieval failed: %s", exc)
        return "", []

    if not results:
        return "", []

    context_parts = []
    sources = []
    for doc, score in results:
        source = doc.metadata.get("source", "unknown")
        context_parts.append(f"[Source: {source}]\n{doc.page_content}")
        sources.append({"source": source, "relevance_score": round(float(score), 3)})

    return "\n\n".join(context_parts), sources


def delete_document_chunks(source_filename: str) -> None:
    """Remove all vector chunks associated with a given source document."""
    store = get_vectorstore()
    try:
        store._collection.delete(where={"source": source_filename})
        logger.info("Deleted vector chunks for '%s'.", source_filename)
    except Exception as exc:
        logger.error("Failed to delete chunks for '%s': %s", source_filename, exc)


def knowledge_base_stats() -> dict:
    """Return basic stats about the current vector store contents."""
    store = get_vectorstore()
    try:
        count = store._collection.count()
    except Exception:
        count = 0
    return {"total_chunks": count}
